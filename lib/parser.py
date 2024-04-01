import os
import time
import uuid

import docx
import moviepy.editor as mp
import cv2 as cv
import numpy as np
import pickle
import whisper
from .constants import *

model = whisper.load_model(WHISPER_MODEL)


class LectureParser:
    def __init__(self, in_video_file, out_document_file, on_status_change=None):
        self.files = []
        self.out_document_file = out_document_file
        self.in_video_file = in_video_file
        self.tmp_audio_file = self.gen_file_name('mp3')
        self.parsed_audio = None
        self.stats = []
        self.is_in_peak = False
        self.fps = None
        self.frames_storage = []
        self.doc = []
        self.on_status_change = on_status_change

    @staticmethod
    def log(self, *args):
        print('LectureParser: ', *args)

    def change_status(self, status):
        """Вызывает callback при изменении статуса"""
        if self.on_status_change:
            self.on_status_change(status)

    def gen_file_name(self, ext=None):
        """Генерирует имя для временного файла"""
        f_name = str(uuid.uuid4()) + ('.' + ext) if ext else ''
        self.files.append(f_name)
        return OUTPUT_FOLDER + f_name

    def remove_files(self):
        """Удаляет временные файлы"""
        for f in self.files:
            if os.path.isfile(OUTPUT_FOLDER + f):
                os.remove(OUTPUT_FOLDER + f)

    def extract_audio_track(self):
        """Извлекает аудио из видео"""
        vc = mp.VideoFileClip(self.in_video_file)
        a = vc.audio
        a.write_audiofile(self.tmp_audio_file, codec='libmp3lame')
        self.log('Extracting audio done')

    def parse_audio(self):
        """Парсит аудио"""
        self.change_status('AUDIO_EXTRACT')
        self.extract_audio_track()
        self.log('Start transcribe')
        self.change_status('AUDIO_TRANSCRIBE')
        result = model.transcribe(self.tmp_audio_file, verbose=True)
        self.log('End transcribe')
        result = list(map(lambda x: [x['start'], x['end'], x['text']], result['segments']))
        self.parsed_audio = result

    def check_peak(self, value):
        if len(self.stats) < PEAK_MIN_COUNT:
            self.stats.append(value)
            return False
        avg = sum(self.stats) / len(self.stats)
        if value > avg * PEAK_RATE:
            return True
        self.stats.append(value)
        if len(self.stats) > PEAK_HIST_COUNT:
            self.stats.pop(0)

    def check_peak_ended(self, value):
        now_in_peak = self.check_peak(value)
        if self.is_in_peak and not now_in_peak:
            self.is_in_peak = now_in_peak
            return True
        self.is_in_peak = now_in_peak
        return False

    def save_snapshot(self, frame_counter, frame):
        f_name = self.gen_file_name('jpeg')
        cv.imwrite(f_name, frame, )
        self.frames_storage.append((frame_counter, f_name))

    def parse_video(self):
        sum_stat = []
        v = cv.VideoCapture(self.in_video_file)
        self.fps = v.get(cv.CAP_PROP_FPS)
        total_frames = int(v.get(cv.CAP_PROP_FRAME_COUNT))
        prev_frame = None
        prev_sum = 0.0
        frame_counter = 0
        kernel = np.ones((5, 5), np.uint8)
        while True:
            ret, frame = v.read()
            frame_counter += 1
            if frame_counter % SKIP_FRAMES != 0:
                continue
            if not ret:
                break
            source_frame = frame
            frame = cv.cvtColor(frame, cv.COLOR_BGR2GRAY)
            frame = cv.blur(frame, (5, 5))
            frame = cv.erode(frame, kernel, iterations=1)
            if TEST_WINDOWS:
                small = cv.resize(frame, (0, 0), fx=0.3, fy=0.3)
                cv.imshow('frame', small)
            if prev_frame is not None:
                diff = cv.absdiff(frame, prev_frame)
                sum_pixels = float(np.sum(diff))
                if TEST_WINDOWS:
                    small = cv.resize(diff, (0, 0), fx=0.3, fy=0.3)
                    cv.imshow('diff', small)
                sum_stat.append([sum_pixels - prev_sum, sum_pixels, prev_sum])
                if self.check_peak_ended(abs(sum_pixels - prev_sum)):
                    success_rate = int(100 * frame_counter / total_frames)
                    self.log('Frame:', success_rate, '%''')
                    self.change_status(f'VIDEO_{success_rate}')
                    self.save_snapshot(frame_counter, source_frame)
                    if TEST_WINDOWS:
                        small = cv.resize(source_frame, (0, 0), fx=0.3, fy=0.3)
                        cv.imshow('snap', small)
                prev_sum = sum_pixels
            if cv.waitKey(1) & 0xFF == ord('q'):
                break
            prev_frame = frame
        # pd.DataFrame(sum_stat).to_csv('sum_stat.csv')

    # def dump_parsed(self):
    #     data = {
    #         'parsed_audio': self.parsed_audio,
    #         'frames_storage': self.frames_storage,
    #         'fps': self.fps
    #     }
    #     with open('output/data_storage.pcl', 'wb') as f:
    #         pickle.dump(data, f)
    #
    # def recover_parsed(self):
    #     with open('output/data_storage.pcl', 'rb') as f:
    #         data = pickle.load(f)
    #         self.parsed_audio = data['parsed_audio']
    #         self.frames_storage = data['frames_storage']
    #         self.fps = data['fps']

    def extract_text(self, from_time_stamp, to_time_stamp):
        return ' '.join(
            map(lambda x: x[2].strip(' '),
                filter(lambda x: from_time_stamp <= x[0] < to_time_stamp, self.parsed_audio)
                )
        )

    def join_video_audio(self):
        self.doc = []
        prev_frame = 0
        for frame_number, f_name in self.frames_storage:
            time_stamp_from = prev_frame / self.fps
            time_stamp_to = frame_number / self.fps
            self.doc.append(
                {'image': f_name, 'text': self.extract_text(time_stamp_from, time_stamp_to)})
            prev_frame = frame_number

    def export_docx(self):
        doc = docx.Document()
        sections = doc.sections
        margin = 1
        for section in sections:
            section.top_margin = docx.shared.Cm(margin)
            section.bottom_margin = docx.shared.Cm(margin)
            section.left_margin = docx.shared.Cm(margin)
            section.right_margin = docx.shared.Cm(margin)
        for item in self.doc:
            doc.add_picture(item['image'], width=docx.shared.Inches(7.5))
            doc.add_paragraph(item['text'].strip(' '))
            self.log('-' * 50)
            self.log(item['image'])
            self.log(item['text'])
        doc.save(self.out_document_file)

    def timer_start(self):
        self.start_time = time.time()

    def timer_print(self, note):
        if SAVE_LOG:
            with open('output/timer_' + self.in_video_file.replace('/', '_').replace('.', '_') + '.log', 'a') as f:
                f.write(note + ' : ' + str(round((time.time() - self.start_time) / 60, 2)) + ' min\n')
        self.log('TIMER: ', note, ' : ', round((time.time() - self.start_time) / 60, 2), ' min')

    def run(self):
        self.timer_start()
        self.change_status('AUDIO')
        self.parse_audio()
        self.timer_print('audio parse')
        self.change_status('VIDEO')
        self.parse_video()
        self.timer_print('video parse')
        self.change_status('MERGE')
        self.join_video_audio()
        self.timer_print('join video and audio')
        self.change_status('EXPORT')
        self.export_docx()
        self.timer_print('export docx')
        self.remove_files()
        self.timer_print('clean up')
        self.timer_print('DONE')
        self.change_status('DONE')


